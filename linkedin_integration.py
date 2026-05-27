#!/usr/bin/env python3
"""
LinkedIn Integration Module for Career Revolution
Extracts LinkedIn URLs from documents and integrates with analysis tools.
"""

import os
import re
import json
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
from docx import Document as DocxDocument

class LinkedInIntegration:
    """Handle LinkedIn URL extraction and profile analysis integration."""
    
    def __init__(self, db_path: str = "career_revolution.db"):
        self.db_path = db_path
        self.linkedin_patterns = [
            r'linkedin\.com/in/[a-zA-Z0-9\-_]+',
            r'linkedin\.com/company/[a-zA-Z0-9\-_]+',
            r'https?://(www\.)?linkedin\.com/[a-zA-Z0-9\-_/]+',
            r'LinkedIn:\s*(https?://[^\s]+)',
            r'Profile:\s*(https?://[^\s]*linkedin[^\s]+)'
        ]
        
    def extract_linkedin_urls_from_text(self, text: str) -> List[str]:
        """Extract LinkedIn URLs from text content."""
        urls = []
        for pattern in self.linkedin_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                url = match.group(0)
                # Clean up the URL - remove "LinkedIn:" prefix if present
                url = re.sub(r'^LinkedIn:\s*', '', url, flags=re.IGNORECASE)
                url = re.sub(r'^Profile:\s*', '', url, flags=re.IGNORECASE)
                
                if not url.startswith('http'):
                    url = 'https://' + url
                # Remove trailing punctuation
                url = url.rstrip('.,;:')
                # Ensure it's a proper LinkedIn URL
                if 'linkedin.com' in url and url not in urls:
                    urls.append(url)
        return urls
    
    def extract_from_pdf(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from PDF file."""
        try:
            text = ""
            urls = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    
                    # Extract URLs from this page
                    page_urls = self.extract_linkedin_urls_from_text(page_text)
                    urls.extend(page_urls)
            
            # Also extract from entire text
            all_urls = self.extract_linkedin_urls_from_text(text)
            unique_urls = list(set(urls + all_urls))
            
            return unique_urls, text[:1000]  # Return first 1000 chars for preview
            
        except Exception as e:
            print(f"Error extracting from PDF {file_path}: {e}")
            return [], ""
    
    def extract_from_docx(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from DOCX file."""
        try:
            text = ""
            urls = []
            
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                text += paragraph_text + "\n"
                
                # Extract URLs from this paragraph
                paragraph_urls = self.extract_linkedin_urls_from_text(paragraph_text)
                urls.extend(paragraph_urls)
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text
                        text += cell_text + "\n"
                        cell_urls = self.extract_linkedin_urls_from_text(cell_text)
                        urls.extend(cell_urls)
            
            # Also extract from entire text
            all_urls = self.extract_linkedin_urls_from_text(text)
            unique_urls = list(set(urls + all_urls))
            
            return unique_urls, text[:1000]  # Return first 1000 chars for preview
            
        except Exception as e:
            print(f"Error extracting from DOCX {file_path}: {e}")
            return [], ""
    
    def extract_from_txt(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            urls = self.extract_linkedin_urls_from_text(text)
            return urls, text[:1000]
            
        except Exception as e:
            print(f"Error extracting from TXT {file_path}: {e}")
            return [], ""
    
    def process_document(self, document_id: int) -> Dict[str, Any]:
        """Process a document from the database and extract LinkedIn URLs."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            # Get document info
            cursor.execute("""
                SELECT id, user_id, original_filename, stored_filename, file_path, 
                       document_type, is_processed
                FROM user_documents 
                WHERE id = ?
            """, (document_id,))
            
            doc = cursor.fetchone()
            if not doc:
                return {"error": "Document not found"}
            
            doc_id, user_id, original_filename, stored_filename, file_path, doc_type, is_processed = doc
            
            # Check if file exists
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}
            
            # Extract based on file type
            linkedin_urls = []
            text_preview = ""
            
            if original_filename.lower().endswith('.pdf'):
                linkedin_urls, text_preview = self.extract_from_pdf(file_path)
            elif original_filename.lower().endswith('.docx'):
                linkedin_urls, text_preview = self.extract_from_docx(file_path)
            elif original_filename.lower().endswith('.txt'):
                linkedin_urls, text_preview = self.extract_from_txt(file_path)
            else:
                return {"error": f"Unsupported file type: {original_filename}"}
            
            # Update document with extracted data
            extracted_data = {
                "linkedin_urls": linkedin_urls,
                "text_preview": text_preview,
                "extraction_date": datetime.now().isoformat(),
                "url_count": len(linkedin_urls)
            }
            
            cursor.execute("""
                UPDATE user_documents 
                SET is_processed = 1,
                    processing_status = 'completed',
                    extracted_data = ?,
                    processed_date = CURRENT_TIMESTAMP
                WHERE id = ?
            """, (json.dumps(extracted_data), doc_id))
            
            # Update user profile with LinkedIn URL if found
            if linkedin_urls:
                # Take the first personal profile URL (not company)
                personal_urls = [url for url in linkedin_urls if '/in/' in url]
                if personal_urls:
                    linkedin_url = personal_urls[0]
                    
                    cursor.execute("""
                        UPDATE user_profiles 
                        SET linkedin_url = ?, updated_at = CURRENT_TIMESTAMP
                        WHERE user_id = ?
                    """, (linkedin_url, user_id))
            
            conn.commit()
            
            # Get user email for response
            cursor.execute("SELECT email FROM users WHERE id = ?", (user_id,))
            user_email = cursor.fetchone()[0]
            
            return {
                "success": True,
                "document_id": doc_id,
                "user_id": user_id,
                "user_email": user_email,
                "original_filename": original_filename,
                "linkedin_urls_found": linkedin_urls,
                "url_count": len(linkedin_urls),
                "text_preview": text_preview[:500] + "..." if len(text_preview) > 500 else text_preview,
                "profile_updated": len(linkedin_urls) > 0
            }
            
        except Exception as e:
            conn.rollback()
            return {"error": f"Processing error: {str(e)}"}
        
        finally:
            conn.close()
    
    def get_user_documents(self, user_id: int) -> List[Dict[str, Any]]:
        """Get all documents for a user with processing status."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, original_filename, document_type, upload_date, 
                   is_processed, processing_status
            FROM user_documents 
            WHERE user_id = ?
            ORDER BY upload_date DESC
        """, (user_id,))
        
        documents = []
        for row in cursor.fetchall():
            doc_id, filename, doc_type, upload_date, is_processed, status = row
            documents.append({
                "id": doc_id,
                "filename": filename,
                "type": doc_type,
                "upload_date": upload_date,
                "is_processed": bool(is_processed),
                "status": status or "pending"
            })
        
        conn.close()
        return documents
    
    def get_user_profile(self, user_id: int) -> Dict[str, Any]:
        """Get user profile with LinkedIn status."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT p.id, p.user_id, u.email, p.linkedin_url, 
                   p.created_at, p.updated_at,
                   (SELECT COUNT(*) FROM user_documents WHERE user_id = p.user_id) as doc_count,
                   (SELECT COUNT(*) FROM user_documents WHERE user_id = p.user_id AND is_processed = 1) as processed_count
            FROM user_profiles p
            JOIN users u ON p.user_id = u.id
            WHERE p.user_id = ?
        """, (user_id,))
        
        row = cursor.fetchone()
        if not row:
            return {}
        
        profile_id, user_id, email, linkedin_url, created_at, updated_at, doc_count, processed_count = row
        
        conn.close()
        
        return {
            "profile_id": profile_id,
            "user_id": user_id,
            "email": email,
            "linkedin_url": linkedin_url,
            "linkedin_status": "found" if linkedin_url else "not_found",
            "documents_total": doc_count,
            "documents_processed": processed_count,
            "created_at": created_at,
            "updated_at": updated_at
        }
    
    def process_all_user_documents(self, user_id: int) -> Dict[str, Any]:
        """Process all unprocessed documents for a user."""
        documents = self.get_user_documents(user_id)
        unprocessed = [doc for doc in documents if not doc["is_processed"]]
        
        results = {
            "user_id": user_id,
            "total_documents": len(documents),
            "unprocessed_documents": len(unprocessed),
            "processed_now": 0,
            "linkedin_urls_found": [],
            "details": []
        }
        
        for doc in unprocessed:
            result = self.process_document(doc["id"])
            if "success" in result and result["success"]:
                results["processed_now"] += 1
                if result["linkedin_urls_found"]:
                    results["linkedin_urls_found"].extend(result["linkedin_urls_found"])
                results["details"].append({
                    "document_id": doc["id"],
                    "filename": doc["filename"],
                    "success": True,
                    "urls_found": result["linkedin_urls_found"]
                })
            else:
                results["details"].append({
                    "document_id": doc["id"],
                    "filename": doc["filename"],
                    "success": False,
                    "error": result.get("error", "Unknown error")
                })
        
        # Remove duplicate URLs
        results["linkedin_urls_found"] = list(set(results["linkedin_urls_found"]))
        
        return results

# Helper function for testing
def test_linkedin_extraction():
    """Test the LinkedIn extraction functionality."""
    print("=" * 60)
    print("TESTING LINKEDIN INTEGRATION MODULE")
    print("=" * 60)
    
    integrator = LinkedInIntegration()
    
    # Test with sample text
    test_text = """
    John Doe
    Software Engineer
    LinkedIn: https://linkedin.com/in/johndoe
    Portfolio: https://github.com/johndoe
    
    Also check my company page: linkedin.com/company/techcorp
    """
    
    print("\n1. Testing URL extraction from text:")
    urls = integrator.extract_linkedin_urls_from_text(test_text)
    print(f"   Found {len(urls)} URLs: {urls}")
    
    # Test database connection
    print("\n2. Testing database connection:")
    try:
        conn = sqlite3.connect("career_revolution.db")
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM users")
        user_count = cursor.fetchone()[0]
        print(f"   Connected to database with {user_count} users")
        conn.close()
    except Exception as e:
        print(f"   Database error: {e}")
    
    # Test getting user documents
    print("\n3. Testing user document retrieval:")
    users = []
    try:
        conn = sqlite3.connect("career_revolution.db")
        cursor = conn.cursor()
        cursor.execute("SELECT id, email FROM users LIMIT 2")
        users = cursor.fetchall()
        conn.close()
    except:
        pass
    
    if users:
        for user_id, email in users:
            print(f"\n   User: {email} (ID: {user_id})")
            profile = integrator.get_user_profile(user_id)
            print(f"   LinkedIn URL: {profile.get('linkedin_url', 'Not found')}")
            documents = integrator.get_user_documents(user_id)
            print(f"   Documents: {len(documents)} total, {sum(1 for d in documents if d['is_processed'])} processed")
    
    print("\n" + "=" * 60)
    print("TEST COMPLETE")
    print("=" * 60)

if __name__ == "__main__":
    test_linkedin_extraction()