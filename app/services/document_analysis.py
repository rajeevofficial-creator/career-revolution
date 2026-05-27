"""
Advanced Document Analysis Service for Career Revolution.
Analyzes uploaded documents to extract career information.
"""

import os
import re
import json
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import PyPDF2
from docx import Document as DocxDocument
import logging
import zipfile
import html

from app.models.schemas import DocumentType
from app.services.linkedin_service import LinkedInService
from app.services.llm_analysis import LLMAnalysisService

logger = logging.getLogger(__name__)

class DocumentAnalysisService:
    """Service for analyzing career documents and extracting structured information."""
    
    def __init__(self):
        self.linkedin_service = LinkedInService()
        self.llm_service = LLMAnalysisService()
        
        # Regex patterns for information extraction
        self.patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(?:\+|00)\d{1,3}[\s.-]?\d{1,3}[\s.-]?\d{2,4}[\s.-]?\d{2,4}(?:[\s.-]?\d{2,4})?|\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            'linkedin': r'(?:https?://)?(?:www\.)?linkedin\.com/(?:in|company)/[a-zA-Z0-9_-]+/?',
            'github': r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+/?',
            'website': r'(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/\S*)?',
            'years_experience': r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?experience',
            'skills': [
                r'Skills?:?\s*([^\.]+(?:\.|$))',
                r'Technical Skills?:?\s*([^\.]+(?:\.|$))',
                r'Programming Languages?:?\s*([^\.]+(?:\.|$))',
                r'Technologies?:?\s*([^\.]+(?:\.|$))'
            ],
            'education': [
                r'Education\s*([^\.]+(?:\.|$))',
                r'Academic Background\s*([^\.]+(?:\.|$))',
                r'Degree[s]?:?\s*([^\.]+(?:\.|$))'
            ],
            'certifications': [
                r'Certifications?:?\s*([^\.]+(?:\.|$))',
                r'Licenses?:?\s*([^\.]+(?:\.|$))',
                r'Credentials?:?\s*([^\.]+(?:\.|$))'
            ],
            'date_range': r'\b(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*|0?\d|1[0-2])[\s/]\d{2,4}\s*(?:-|to|–|until)\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*|0?\d|1[0-2]|Present|Current)[\s/]\d{0,4}\b',
            'year_range': r'\b(19|20)\d{2}\s*[-–]\s*(?:(19|20)\d{2}|Present|Current)\b'
        }
        
        # Common skills dictionary for extraction
        self.common_skills = {
            'programming': ['python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring', 'express'],
            'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible'],
            'data_science': ['pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn', 'r', 'matplotlib'],
            'devops': ['jenkins', 'git', 'github', 'gitlab', 'ci/cd', 'linux', 'bash'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'project management']
        }
        
        # Document type detection patterns - using string values since enum has limited options
        self.document_type_patterns = {
            "resume": [
                r'\b(?:resume|curriculum vitae|cv)\b',
                r'\b(?:work experience|professional experience)\b',
                r'\b(?:summary|objective)\b',
                r'\b(?:skills|technical skills)\b'
            ],
            "certification": [
                r'\b(?:certificate|certification)\b',
                r'\b(?:issued by|certified by)\b',
                r'\b(?:valid until|expiration date)\b',
                r'\b(?:license number|certification id)\b'
            ],
            "portfolio": [
                r'\b(?:portfolio|projects)\b',
                r'\b(?:github|gitlab|bitbucket)\b',
                r'\b(?:live demo|deployed at)\b',
                r'\b(?:technologies used|built with)\b'
            ],
            "experience_letter": [
                r'\b(?:experience letter|employment certificate)\b',
                r'\b(?:to whom it may concern)\b',
                r'\b(?:this is to certify that)\b',
                r'\b(?:worked as|served as)\b'
            ],
            "transcript": [
                r'\b(?:transcript|academic record)\b',
                r'\b(?:grade|gpa|cgpa)\b',
                r'\b(?:course|subject|credit)\b',
                r'\b(?:semester|term)\b'
            ]
        }
        
        self.linkedin_service = LinkedInService()
    
    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Robustly parse dates from document text."""
        if not date_str or not isinstance(date_str, str):
            return None
            
        date_str = date_str.strip().lower()
        if 'present' in date_str or 'current' in date_str or 'now' in date_str:
            # We use None or special value for Present in DB, but for start_date we need a real date
            return datetime.utcnow()
            
        # Try various formats
        formats = [
            '%Y-%m', '%m/%Y', '%b %Y', '%B %Y', '%m-%Y', '%Y',
            '%b %y', '%B %y', '%m/%y'
        ]
        
        # Clean string from common noise but keep / - for formats
        clean_str = re.sub(r'[^a-zA-Z0-9/\-]', ' ', date_str).strip()
        
        for fmt in formats:
            try:
                dt = datetime.strptime(clean_str, fmt)
                return dt
            except ValueError:
                continue
                
        # Try regex for just year
        year_match = re.search(r'\b(19|20)\d{2}\b', clean_str)
        if year_match:
            try:
                return datetime(year=int(year_match.group(0)), month=1, day=1)
            except ValueError:
                pass
                
        return None

    def extract_text_from_file(self, file_path: str, file_ext: str) -> Optional[str]:
        """Extract text from different file types."""
        try:
            content = ""
            if file_ext.lower() == 'pdf':
                content = self._extract_text_from_pdf(file_path)
            elif file_ext.lower() in ['doc', 'docx']:
                content = self._extract_text_from_docx(file_path)
            elif file_ext.lower() == 'txt':
                content = self._extract_text_from_txt(file_path)
            elif file_ext.lower() in ['png', 'jpg', 'jpeg']:
                content = self._extract_text_from_image(file_path)
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_ext}")
                return ""
            
            # If text is too short or likely garbage, try OCR fallback for PDF/Images
            if len(content.strip()) < 100 and file_ext.lower() in ['pdf', 'png', 'jpg', 'jpeg']:
                logger.info(f"Low text yield for {file_path}, attempting OCR fallback.")
                ocr_text = self._extract_text_via_ocr(file_path, file_ext)
                if len(ocr_text) > len(content):
                    return ocr_text
            
            return content
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image files via OCR."""
        return self._extract_text_via_ocr(file_path, 'image')

    def _extract_text_via_ocr(self, file_path: str, file_type: str) -> str:
        """Generic OCR extraction using pytesseract."""
        try:
            import pytesseract
            from PIL import Image
            
            if file_type.lower() == 'pdf':
                from pdf2image import convert_from_path
                # Convert first 3 pages to images for OCR
                images = convert_from_path(file_path, last_page=3)
                text = ""
                for img in images:
                    text += pytesseract.image_to_string(img) + "\n"
                return text
            else:
                return pytesseract.image_to_string(Image.open(file_path))
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return ""
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF (fitz) with PyPDF2 fallback."""
        text = ""
        try:
            import fitz  # PyMuPDF
            with fitz.open(file_path) as doc:
                for page in doc:
                    extracted = page.get_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path} with PyMuPDF: {e}")
            
        # Fallback to PyPDF2 if fitz fails or yields very little text
        if len(text.strip()) < 50:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
            except Exception as e2:
                logger.error(f"Error reading PDF {file_path} with PyPDF2 fallback: {e2}")
                
        # Second fallback: pdfplumber
        if len(text.strip()) < 50:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as doc:
                    for page in doc.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
            except Exception as e3:
                logger.error(f"Error reading PDF {file_path} with pdfplumber fallback: {e3}")
                
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file including tables and XML fallback."""
        text_parts = []
        try:
            doc = DocxDocument(file_path)
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            # If standard extraction returns very little, try XML fallback
            if len(text.strip()) < 100:
                logger.info(f"Standard DOCX extraction failed to find content for {file_path}. Trying XML fallback.")
                xml_text = self._extract_text_from_docx_xml(file_path)
                if len(xml_text) > len(text):
                    return xml_text
            
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            # Try XML fallback on error
            return self._extract_text_from_docx_xml(file_path)

    def _extract_text_from_docx_xml(self, file_path: str) -> str:
        """Robust fallback to extract text from DOCX XML directly."""
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                if 'word/document.xml' not in z.namelist():
                    return ""
                
                with z.open('word/document.xml') as doc_xml:
                    xml_content = doc_xml.read().decode('utf-8', errors='ignore')
                    
                    # 1. Replace paragraph ends and breaks with newlines
                    content = re.sub(r'</w:p>', '\n', xml_content)
                    content = re.sub(r'<w:br[^>]*>', '\n', content)
                    # 2. Strip all other tags
                    content = re.sub(r'<[^>]+>', '', content)
                    # 3. Collapse multiple spaces and newlines
                    content = re.sub(r' +', ' ', content)
                    content = re.sub(r'\n+', '\n', content)
                    
                    return html.unescape(content).strip()
        except Exception as e:
            logger.error(f"Error in DOCX XML fallback for {file_path}: {e}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def detect_document_type(self, filename: str, content: str) -> str:
        """Detect the type of document based on filename and content."""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        # Check filename patterns first
        if any(pattern in filename_lower for pattern in ['resume', 'cv', 'curriculum']):
            return "resume"
        elif any(pattern in filename_lower for pattern in ['certificate', 'certification', 'license']):
            return "certification"
        elif any(pattern in filename_lower for pattern in ['experience', 'employment', 'letter']):
            return "portfolio"  # Using portfolio as fallback for experience letters
        elif any(pattern in filename_lower for pattern in ['transcript', 'grades', 'academic']):
            return "portfolio"  # Using portfolio as fallback for transcripts
        elif any(pattern in filename_lower for pattern in ['portfolio', 'projects', 'github']):
            return "portfolio"
        
        # Check content patterns
        for doc_type, patterns in self.document_type_patterns.items():
            for pattern in patterns:
                if re.search(pattern, content_lower, re.IGNORECASE):
                    return doc_type
        
        # Default to "resume" if no specific type detected
        return "resume"
    
    def extract_personal_info(self, content: str) -> Dict[str, Any]:
        """Extract personal information from document content."""
        info = {
            'emails': [],
            'phones': [],
            'linkedin_urls': [],
            'github_urls': [],
            'websites': [],
            'name': None,
            'location': None,
            'summary': None
        }
        
        # 1. Regex extractions for structured data
        info['emails'] = list(set(re.findall(self.patterns['email'], content, re.IGNORECASE)))
        info['phones'] = list(set(re.findall(self.patterns['phone'], content, re.IGNORECASE)))
        info['linkedin_urls'] = list(set(re.findall(self.patterns['linkedin'], content, re.IGNORECASE)))
        info['github_urls'] = list(set(re.findall(self.patterns['github'], content, re.IGNORECASE)))
        info['websites'] = list(set(re.findall(self.patterns['website'], content, re.IGNORECASE)))

        # 2. Try to extract name
        lines = content.split('\n')
        # Skip common section headers that might be capitalized
        headers = ['resume', 'cv', 'curriculum vitae', 'summary', 'contact', 'experience', 'education', 'skills']
        
        for line in lines[:50]:
            line = line.strip()
            if not line or len(line) < 3: continue
            if '@' in line or 'http' in line or re.search(r'\d{5,}', line): continue
            
            if line.lower() in headers: continue
            
            # Pattern: 2-4 words, each starting with capital
            if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z](?:[a-z]+|\.)){1,3}$', line):
                info['name'] = line
                break
        
        # 3. Location extraction - look for "City, Country" or similar
        location_pattern = r'\b([A-Z][a-z]+(?:[\s-][A-Z][a-z]+)*,\s*[A-Z][a-z]+(?:[\s-][A-Z][a-z]+)*|[A-Z][a-z]+,\s*[A-Z]{2})\b'
        loc_match = re.search(location_pattern, content)
        if loc_match:
            info['location'] = loc_match.group(1)

        # 4. Extract summary (look for "SUMMARY" or "PROFILE" headers)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        summary_sections = ['summary', 'profile', 'professional summary', 'objective', 'about me']
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()
            match = None
            for s in summary_sections:
                if para_lower.startswith(s):
                    match = s
                    break
            
            if match:
                summary_text = para[len(match):].strip().lstrip(':').strip()
                if not summary_text and i + 1 < len(paragraphs):
                    summary_text = paragraphs[i+1]
                
                # Filter boilerplate
                if len(summary_text) > 20 and "this is to certify" not in summary_text.lower():
                    info['summary'] = summary_text
                break
        
        # Fallback summary: first substantial paragraph that isn't contact info or boilerplate
        if not info['summary']:
            for para in paragraphs[:5]:
                if len(para) > 60 and '@' not in para and 'http' not in para:
                    if para.count(',') < 8 and "this is to certify" not in para.lower():
                        info['summary'] = para
                        break
        
        return info
    
    def extract_skills(self, content: str) -> List[str]:
        """Extract skills from document content."""
        skills = []
        content_lower = content.lower()
        
        # Extract using regex patterns
        # More specific skill extraction patterns to avoid multi-line capture
        skill_patterns = [
            r'Skills?:?\s*([^\.\n]+(?:\.|$|\n))',
            r'Technical Skills?:?\s*([^\.\n]+(?:\.|$|\n))',
            r'Programming Languages?:?\s*([^\.\n]+(?:\.|$|\n))',
            r'Technologies?:?\s*([^\.\n]+(?:\.|$|\n))'
        ]
        
        for pattern in skill_patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE)
            for match in matches:
                skill_text = match.group(1).strip()
                
                # Split by common delimiters
                if skill_text.endswith('.'):
                    skill_text = skill_text[:-1]
                
                # Normalize delimiters
                for delimiter in [',', ';', '|', '/', '•', '-', '–']:
                    skill_text = skill_text.replace(delimiter, ',')
                
                # Extract individual skills (strict length and noise filter)
                potential_skills = [
                    s.strip() for s in skill_text.split(',') 
                    if s.strip() and 2 < len(s.strip()) < 30 and '\n' not in s.strip()
                ]
                skills.extend(potential_skills)
        
        # Also check for common skills
        for category, skill_list in self.common_skills.items():
            for skill in skill_list:
                if skill in content_lower:
                    skills.append(skill)
        
        # Clean and deduplicate
        skills = list(set([s.strip().title() for s in skills if s.strip()]))
        
        return skills
    
    def extract_education(self, content: str) -> List[Dict[str, Any]]:
        """Extract education information."""
        education = []
        
        # Common noise to exclude
        noise = [
            'education', 'academic', 'certifications', 'licenses', 
            'credentials', 'skills', 'experience', 'summary', 
            'contact', 'profile', 'objective'
        ]
        
        for pattern in self.patterns['education']:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                # Simple extraction - could be enhanced with NLP
                text_block = match.group(1)
                lines = [line.strip() for line in text_block.split('\n') if line.strip()]
                for line in lines:
                    line_lower = line.lower().strip(': ')
                    if any(n == line_lower for n in noise):
                        continue
                        
                    if len(line) > 10:  # Reasonable length for education entry
                        # Check if it looks like a degree or institution
                        education.append({
                            'institution': line,
                            'degree': "Degree",  # Placeholder
                            'years': None,
                            'description': line
                        })
        
        return education
    
    def extract_certifications(self, content: str) -> List[Dict[str, str]]:
        """Extract certifications."""
        certifications = []
        
        for pattern in self.patterns['certifications']:
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                lines = [line.strip() for line in match.split('\n') if line.strip()]
                for line in lines:
                    if len(line) > 5:  # Reasonable length for certification
                        certifications.append({
                            'name': line,
                            'issuer': None,  # Could be extracted
                            'date': None,
                            'description': line
                        })
        
        return certifications
    
    def extract_experience(self, content: str) -> List[Dict[str, str]]:
        """Extract work experience information with improved logic using date ranges."""
        experience = []
        
        # Find all date ranges - these are often anchors for job entries
        date_ranges = list(re.finditer(self.patterns['date_range'], content, re.IGNORECASE))
        if not date_ranges:
            date_ranges = list(re.finditer(self.patterns['year_range'], content, re.IGNORECASE))
            
        if date_ranges:
            # Use date ranges as anchors to find job titles and companies
            for i, match in enumerate(date_ranges):
                start_pos = match.start()
                end_pos = match.end()
                
                # Look at 100 characters before the date range for title/company
                # Skip empty lines to find the actual content line
                anchor_lines = [l.strip() for l in content[max(0, start_pos-150):start_pos].split('\n') if l.strip()]
                relevant_line = anchor_lines[-1] if anchor_lines else ""
                
                # If the line is just the date or part of a date, try the line before it
                if len(relevant_line) < 5 and len(anchor_lines) > 1:
                    relevant_line = anchor_lines[-2]
                
                # Heuristic for title/company: "Job Title at Company" or "Company | Job Title" or "Company - Job Title"
                title = None
                company = None
                
                if ' at ' in relevant_line.lower():
                    parts = re.split(r'\s+at\s+', relevant_line, flags=re.IGNORECASE)
                    title = parts[0].strip() or "Professional"
                    company = parts[1].strip() or "Unknown Company"
                elif ' | ' in relevant_line:
                    parts = relevant_line.split(' | ')
                    # Heuristic: usually titles contain certain keywords
                    pos_keywords = ['engineer', 'developer', 'manager', 'lead', 'analyst', 'designer', 'director', 'consultant']
                    if any(kw in parts[0].lower() for kw in pos_keywords):
                        title, company = parts[0].strip(), parts[1].strip()
                    else:
                        company, title = parts[0].strip(), parts[1].strip()
                elif ' - ' in relevant_line:
                    parts = relevant_line.split(' - ')
                    if any(kw in parts[0].lower() for kw in ['engineer', 'developer', 'manager', 'lead']):
                        title, company = parts[0].strip(), parts[1].strip()
                    else:
                        company, title = parts[0].strip(), parts[1].strip()
                else:
                    # Single line: Could be Title, Company, or Title - Company
                    title = relevant_line or "Professional"
                    company = "Unknown Company"
                
                # Look ahead for description
                next_anchor = date_ranges[i+1].start() if i + 1 < len(date_ranges) else len(content)
                description_text = content[end_pos:next_anchor].strip()
                # Take the first few lines as description
                description_lines = [l.strip() for l in description_text.split('\n') if len(l.strip()) > 15]
                description = ' '.join(description_lines[:3])
                
                if title and len(title) > 3:
                    experience.append({
                        'title': title,
                        'company': company,
                        'start_date': self._parse_date(match.group(1).split('-')[0].strip()) if '-' in match.group(0) else None,
                        'end_date': self._parse_date(match.group(1).split('-')[1].strip()) if '-' in match.group(0) and len(match.group(1).split('-')) > 1 else None,
                        'duration': match.group(0),
                        'description': description
                    })
        
        # Fallback to simple line-based extraction if no date ranges found
        if not experience:
            lines = content.split('\n')
            current_job = None
            for i, line in enumerate(lines):
                line = line.strip()
                if re.search(r'\b(?:Senior|Junior|Lead|Principal|Manager|Director|Engineer|Developer|Analyst|Consultant)\b', line, re.IGNORECASE):
                    if len(line) < 100:
                        current_job = {'title': line, 'company': None, 'duration': None, 'description': []}
                elif current_job and not current_job['company'] and re.search(r'\b(?:at|@|for)\b', line, re.IGNORECASE):
                    current_job['company'] = line
                elif current_job and line and len(line) > 20:
                    current_job['description'].append(line)
                elif current_job and (not line or i == len(lines) - 1):
                    if current_job['description']:
                        current_job['description'] = ' '.join(current_job['description'][:3])
                        experience.append(current_job)
                    current_job = None
        
        return experience
    
    async def analyze_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Comprehensive analysis of a single document."""
        logger.info(f"Analyzing document: {filename}")
        
        # Get file extension
        file_ext = os.path.splitext(filename)[1].lower().replace('.', '')
        
        # Extract text first (local)
        content = self.extract_text_from_file(file_path, file_ext)
        
        if not content and file_ext not in ['png', 'jpg', 'jpeg', 'pdf']:
            return {
                'filename': filename,
                'error': 'Could not extract text from document',
                'analysis_complete': False
            }
        
        # Determine analysis mode
        analysis_result = {}
        
        # If LLM is available, use it for higher fidelity
        if self.llm_service.is_available():
            logger.info(f"Using LLM for high-fidelity analysis of {filename}")
            if file_ext in ['png', 'jpg', 'jpeg']:
                with open(file_path, "rb") as f:
                    image_bytes = f.read()
                mime_type = f"image/{file_ext if file_ext != 'jpg' else 'jpeg'}"
                llm_result = await self.llm_service.analyze_file(image_bytes, mime_type)
            elif file_ext == 'pdf' and not content:
                # Scanned PDF: Use vision analysis
                logger.info(f"Scanned PDF detected for {filename}, using vision analysis.")
                with open(file_path, "rb") as f:
                    pdf_bytes = f.read()
                llm_result = await self.llm_service.analyze_file(pdf_bytes, "application/pdf")
            else:
                llm_result = await self.llm_service.analyze_text(content, document_context=filename)
            
            if "error" not in llm_result:
                # Merge LLM results with basic metadata
                analysis_result = {
                    'filename': filename,
                    'document_type': self.detect_document_type(filename, content),
                    'analysis_complete': True,
                    'analysis_timestamp': datetime.utcnow().isoformat(),
                    'source': 'llm'
                }
                analysis_result.update(llm_result)
                
                # HYBRID MERGE: Re-run heuristic extraction for contact info (regex-based is highly reliable)
                heuristic_info = self.extract_personal_info(content)
                heuristic_linkedin = self.linkedin_service.extract_linkedin_urls_from_text(content)
                
                if 'personal_info' not in analysis_result:
                    analysis_result['personal_info'] = {}
                
                pi = analysis_result['personal_info']
                
                # Merge LinkedIn
                if heuristic_linkedin:
                    analysis_result['linkedin_urls'] = list(set(analysis_result.get('linkedin_urls', []) + heuristic_linkedin))
                    if not pi.get('linkedin_url'):
                        p_urls = [url for url in heuristic_linkedin if '/in/' in url]
                        if p_urls: pi['linkedin_url'] = p_urls[0]
                
                # Merge other contact info if LLM missed it
                if not pi.get('phone') and heuristic_info.get('phones'):
                    pi['phone'] = heuristic_info['phones'][0]
                if not pi.get('email') and heuristic_info.get('emails'):
                    pi['email'] = heuristic_info['emails'][0]
                if not pi.get('location') and heuristic_info.get('location'):
                    pi['location'] = heuristic_info['location']
                if not pi.get('name') and heuristic_info.get('name'):
                    pi['name'] = heuristic_info['name']

                # Ensure fields exist
                for field in ['skills', 'experience', 'education', 'certifications', 'personal_info']:
                    if field not in analysis_result:
                        analysis_result[field] = [] if field != 'personal_info' else {}
                
                return analysis_result
            else:
                logger.warning(f"LLM analysis failed for {filename}: {llm_result['error']}. Falling back to heuristics.")

        # Heuristic fallback (original logic)
        doc_type = self.detect_document_type(filename, content)
        personal_info = self.extract_personal_info(content)
        skills = self.extract_skills(content)
        education = self.extract_education(content)
        certifications = self.extract_certifications(content)
        experience = self.extract_experience(content)
        linkedin_urls = self.linkedin_service.extract_linkedin_urls_from_text(content)
        
        years_match = re.search(self.patterns['years_experience'], content, re.IGNORECASE)
        years_experience = int(years_match.group(1)) if years_match else None
        
        analysis_result = {
            'filename': filename,
            'document_type': doc_type,
            'content_preview': content[:500] + '...' if len(content) > 500 else content,
            'personal_info': personal_info,
            'skills': skills,
            'education': education[:5],
            'certifications': certifications[:10],
            'experience': experience[:5],
            'linkedin_urls': linkedin_urls,
            'years_experience': years_experience,
            'analysis_timestamp': datetime.utcnow().isoformat(),
            'analysis_complete': True,
            'source': 'heuristic'
        }
        
        return analysis_result
    
    def analyze_folder(self, folder_path: str) -> Dict[str, Any]:
        """Analyze all documents in a folder."""
        logger.info(f"Analyzing folder: {folder_path}")
        
        if not os.path.exists(folder_path):
            return {
                'error': f'Folder not found: {folder_path}',
                'analysis_complete': False
            }
        
        # Get all files in folder
        all_files = []
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                all_files.append(file_path)
        
        # Filter for supported file types
        supported_extensions = ['.pdf', '.doc', '.docx', '.txt', '.png', '.jpg', '.jpeg']
        supported_files = [
            f for f in all_files 
            if os.path.splitext(f)[1].lower() in supported_extensions
        ]
        
        if not supported_files:
            return {
                'error': 'No supported files found in folder',
                'analysis_complete': False
            }
        
        # Analyze each file
        analysis_results = []
        for file_path in supported_files:
            filename = os.path.basename(file_path)
            try:
                result = self.analyze_document(file_path, filename)
                analysis_results.append(result)
            except Exception as e:
                logger.error(f"Error analyzing {filename}: {e}")
                analysis_results.append({
                    'filename': filename,
                    'error': str(e),
                    'analysis_complete': False
                })
        
        # Aggregate results
        aggregated = self._aggregate_analysis_results(analysis_results)
        
        logger.info(f"Folder analysis complete: {len(analysis_results)} files analyzed")
        
        return aggregated
    
    def _aggregate_analysis_results(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Aggregate analysis results from multiple documents."""
        aggregated = {
            'total_files': len(results),
            'successful_analyses': sum(1 for r in results if r.get('analysis_complete', False)),
            'failed_analyses': sum(1 for r in results if not r.get('analysis_complete', False)),
            'document_types': {},
            'all_skills': [],
            'all_education': [],
            'all_certifications': [],
            'all_experience': [],
            'all_linkedin_urls': [],
            'personal_info_consolidated': {
                'emails': set(),
                'phones': set(),
                'linkedin_urls': set(),
                'github_urls': set(),
                'websites': set(),
                'names': set()
            },
            'individual_results': results,
            'summary': {}
        }
        
        # Process each result
        for result in results:
            if not result.get('analysis_complete', False):
                continue
            
            # Count document types
            doc_type = result.get('document_type', 'unknown')
            aggregated['document_types'][doc_type] = aggregated['document_types'].get(doc_type, 0) + 1
            
            # Collect skills
            skills = result.get('skills', [])
            aggregated['all_skills'].extend(skills)
            
            # Collect education
            education = result.get('education', [])
            aggregated['all_education'].extend(education)
            
            # Collect certifications
            certifications = result.get('certifications', [])
            aggregated['all_certifications'].extend(certifications)
            
            # Collect experience
            experience = result.get('experience', [])
            aggregated['all_experience'].extend(experience)
            
            # Collect LinkedIn URLs
            linkedin_urls = result.get('linkedin_urls', [])
            aggregated['all_linkedin_urls'].extend(linkedin_urls)
            
            # Consolidate personal info
            personal_info = result.get('personal_info', {})
            for key in ['emails', 'phones', 'linkedin_urls', 'github_urls', 'websites']:
                if key in personal_info:
                    aggregated['personal_info_consolidated'][key].update(personal_info[key])
            
            if personal_info.get('name'):
                aggregated['personal_info_consolidated']['names'].add(personal_info['name'])
        
        # Convert sets to lists for JSON serialization
        for key in aggregated['personal_info_consolidated']:
            aggregated['personal_info_consolidated'][key] = list(aggregated['personal_info_consolidated'][key])
        
        # Deduplicate and sort
        aggregated['all_skills'] = sorted(list(set(aggregated['all_skills'])))
        aggregated['all_linkedin_urls'] = sorted(list(set(aggregated['all_linkedin_urls'])))
        
        # Create summary
        aggregated['summary'] = {
            'total_skills_identified': len(aggregated['all_skills']),
            'total_certifications': len(aggregated['all_certifications']),
            'total_education_entries': len(aggregated['all_education']),
            'total_experience_entries': len(aggregated['all_experience']),
            'primary_document_type': max(aggregated['document_types'].items(), key=lambda x: x[1])[0] if aggregated['document_types'] else 'unknown',
            'has_linkedin_urls': len(aggregated['all_linkedin_urls']) > 0,
            'has_contact_info': any(len(v) > 0 for k, v in aggregated['personal_info_consolidated'].items() if k != 'names')
        }
        
        return aggregated
    
    def save_analysis_to_profile(self, db, user_id: int, analysis_results: Dict[str, Any]) -> Dict[str, Any]:
        """Save analysis results to user profile."""
        # This would integrate with the profile service
        # For now, return the results
        return {
            'message': 'Analysis results ready for review',
            'analysis_id': f"analysis_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}",
            'results': analysis_results,
            'profile_update_suggestions': self._generate_profile_suggestions(analysis_results)
        }
    
    def _generate_profile_suggestions(self, analysis_results: Dict[str, Any]) -> Dict[str, Any]:
        """Generate suggestions for profile updates based on analysis."""
        suggestions = {
            'skills_to_add': [],
            'education_to_add': [],
            'certifications_to_add': [],
            'experience_to_add': [],
            'linkedin_url_to_add': None,
            'summary_suggestion': None
        }
        
        # Suggest adding top skills
        if analysis_results.get('all_skills'):
            suggestions['skills_to_add'] = analysis_results['all_skills'][:10]  # Top 10 skills
        
        # Suggest adding education
        if analysis_results.get('all_education'):
            suggestions['education_to_add'] = analysis_results['all_education'][:3]  # Top 3 education entries
        
        # Suggest adding certifications
        if analysis_results.get('all_certifications'):
            suggestions['certifications_to_add'] = analysis_results['all_certifications'][:5]  # Top 5 certifications
        
        # Suggest adding experience
        if analysis_results.get('all_experience'):
            suggestions['experience_to_add'] = analysis_results['all_experience'][:3]  # Top 3 experience entries
        
        # Suggest LinkedIn URL
        if analysis_results.get('all_linkedin_urls'):
            # Prefer personal profile URLs over company URLs
            personal_urls = [url for url in analysis_results['all_linkedin_urls'] if '/in/' in url]
            if personal_urls:
                suggestions['linkedin_url_to_add'] = personal_urls[0]
            else:
                suggestions['linkedin_url_to_add'] = analysis_results['all_linkedin_urls'][0]
        
        # Generate summary suggestion
        total_skills = len(analysis_results.get('all_skills', []))
        total_exp = analysis_results.get('summary', {}).get('total_experience_entries', 0)
        
        if total_skills > 0 and total_exp > 0:
            suggestions['summary_suggestion'] = (
                f"Professional with {total_exp} key experience entries and expertise in {total_skills} skills "
                f"including {', '.join(analysis_results.get('all_skills', [])[:3])}."
            )
        
        return suggestions
