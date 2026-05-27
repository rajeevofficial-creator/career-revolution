"""
LinkedIn integration service for Career Revolution.
Handles LinkedIn URL extraction from documents and profile analysis.
"""

import os
import re
import json
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from sqlalchemy.orm import Session
import PyPDF2
from docx import Document as DocxDocument

from app.models.database import UserDocument, UserProfile
from app.models.schemas import LinkedInExtractionResult, LinkedInAnalysisOptions

class LinkedInService:
    """Service for LinkedIn URL extraction and profile analysis."""
    
    def __init__(self):
        self.linkedin_patterns = [
            r'(?:[a-z]{2}\.)?linkedin\.com/in/[a-zA-Z0-9\-_]+/?',
            r'(?:[a-z]{2}\.)?linkedin\.com/company/[a-zA-Z0-9\-_]+/?',
            r'https?://(?:[a-z]{2}\.)?linkedin\.com/[a-zA-Z0-9\-_/]+',
            r'LinkedIn:\s*((?:https?://)?[^\s]+)',
            r'Profile:\s*((?:https?://)?[^\s]*linkedin[^\s]+)',
            r'LinkedIn Profile:\s*((?:https?://)?[^\s]+)'
        ]
    
    def extract_linkedin_urls_from_text(self, text: str) -> List[str]:
        """Extract LinkedIn URLs from text content."""
        urls = []
        for pattern in self.linkedin_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                url = match.group(0)
                # Clean up the URL - remove prefixes
                url = re.sub(r'^LinkedIn:\s*', '', url, flags=re.IGNORECASE)
                url = re.sub(r'^Profile:\s*', '', url, flags=re.IGNORECASE)
                url = re.sub(r'^LinkedIn Profile:\s*', '', url, flags=re.IGNORECASE)
                url = re.sub(r'^Connect with me on LinkedIn:\s*', '', url, flags=re.IGNORECASE)
                
                if not url.startswith('http'):
                    url = 'https://' + url
                # Remove trailing punctuation
                url = url.rstrip('.,;:')
                # Ensure it's a proper LinkedIn URL
                if 'linkedin.com' in url and url not in urls:
                    urls.append(url)
        return urls
    
    def extract_from_pdf(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from PDF file."""
        try:
            text = ""
            urls = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    
                    # Extract URLs from this page
                    page_urls = self.extract_linkedin_urls_from_text(page_text)
                    urls.extend(page_urls)
            
            # Also extract from entire text
            all_urls = self.extract_linkedin_urls_from_text(text)
            unique_urls = list(set(urls + all_urls))
            
            return unique_urls, text[:1000]  # Return first 1000 chars for preview
            
        except Exception as e:
            print(f"Error extracting from PDF {file_path}: {e}")
            return [], ""
    
    def extract_from_docx(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from DOCX file."""
        try:
            text = ""
            urls = []
            
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                text += paragraph_text + "\n"
                
                # Extract URLs from this paragraph
                paragraph_urls = self.extract_linkedin_urls_from_text(paragraph_text)
                urls.extend(paragraph_urls)
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text
                        text += cell_text + "\n"
                        cell_urls = self.extract_linkedin_urls_from_text(cell_text)
                        urls.extend(cell_urls)
            
            # Also extract from entire text
            all_urls = self.extract_linkedin_urls_from_text(text)
            unique_urls = list(set(urls + all_urls))
            
            return unique_urls, text[:1000]  # Return first 1000 chars for preview
            
        except Exception as e:
            print(f"Error extracting from DOCX {file_path}: {e}")
            return [], ""
    
    def extract_from_txt(self, file_path: str) -> Tuple[List[str], str]:
        """Extract text and LinkedIn URLs from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            urls = self.extract_linkedin_urls_from_text(text)
            return urls, text[:1000]
            
        except Exception as e:
            print(f"Error extracting from TXT {file_path}: {e}")
            return [], ""
    
    def process_document_for_linkedin(self, db: Session, document_id: int) -> Dict[str, Any]:
        """Process a document to extract LinkedIn URLs."""
        # Get document from database
        document = db.query(UserDocument).filter(UserDocument.id == document_id).first()
        if not document:
            return {"error": "Document not found"}
        
        # Check if file exists
        if not os.path.exists(document.file_path):
            return {"error": f"File not found: {document.file_path}"}
        
        # Extract based on file type
        linkedin_urls = []
        text_preview = ""
        
        original_filename = document.original_filename.lower()
        
        if original_filename.endswith('.pdf'):
            linkedin_urls, text_preview = self.extract_from_pdf(document.file_path)
        elif original_filename.endswith('.docx'):
            linkedin_urls, text_preview = self.extract_from_docx(document.file_path)
        elif original_filename.endswith('.txt'):
            linkedin_urls, text_preview = self.extract_from_txt(document.file_path)
        else:
            return {"error": f"Unsupported file type: {document.original_filename}"}
        
        # Update document with extracted data
        extracted_data = {
            "linkedin_urls": linkedin_urls,
            "text_preview": text_preview,
            "extraction_date": datetime.now().isoformat(),
            "url_count": len(linkedin_urls)
        }
        
        document.is_processed = True
        document.processing_status = "completed"
        document.extracted_data = json.dumps(extracted_data)
        document.processed_date = datetime.now()
        
        # Update user profile with LinkedIn URL if found
        if linkedin_urls:
            # Take the first personal profile URL (not company)
            personal_urls = [url for url in linkedin_urls if '/in/' in url]
            if personal_urls:
                linkedin_url = personal_urls[0]
                
                # Get or create user profile
                profile = db.query(UserProfile).filter(UserProfile.user_id == document.user_id).first()
                if not profile:
                    profile = UserProfile(user_id=document.user_id)
                    db.add(profile)
                
                profile.linkedin_url = linkedin_url
                profile.updated_at = datetime.now()
        
        db.commit()
        
        return {
            "success": True,
            "document_id": document.id,
            "user_id": document.user_id,
            "original_filename": document.original_filename,
            "linkedin_urls_found": linkedin_urls,
            "url_count": len(linkedin_urls),
            "text_preview": text_preview[:500] + "..." if len(text_preview) > 500 else text_preview,
            "profile_updated": len(linkedin_urls) > 0
        }
    
    def get_user_linkedin_status(self, db: Session, user_id: int) -> Dict[str, Any]:
        """Get LinkedIn status for a user."""
        # Get user profile
        profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
        
        # Get user documents
        documents = db.query(UserDocument).filter(
            UserDocument.user_id == user_id
        ).order_by(UserDocument.upload_date.desc()).all()
        
        # Count processed documents
        total_docs = len(documents)
        processed_docs = sum(1 for doc in documents if doc.is_processed)
        
        # Check if any documents have LinkedIn URLs
        linkedin_urls = []
        for doc in documents:
            if doc.extracted_data:
                try:
                    data = json.loads(doc.extracted_data)
                    if "linkedin_urls" in data:
                        linkedin_urls.extend(data["linkedin_urls"])
                except:
                    pass
        
        # Remove duplicates
        linkedin_urls = list(set(linkedin_urls))
        
        return {
            "user_id": user_id,
            "linkedin_url": profile.linkedin_url if profile else None,
            "linkedin_status": "found" if profile and profile.linkedin_url else "not_found",
            "documents_total": total_docs,
            "documents_processed": processed_docs,
            "linkedin_urls_found": linkedin_urls,
            "profile_exists": profile is not None,
            "profile_updated": profile.updated_at if profile else None
        }
    
    def process_all_user_documents(self, db: Session, user_id: int) -> Dict[str, Any]:
        """Process all unprocessed documents for a user."""
        # Get all unprocessed documents for the user
        documents = db.query(UserDocument).filter(
            UserDocument.user_id == user_id,
            UserDocument.is_processed == False
        ).all()
        
        results = {
            "user_id": user_id,
            "total_documents": len(documents),
            "processed_now": 0,
            "linkedin_urls_found": [],
            "details": []
        }
        
        for document in documents:
            result = self.process_document_for_linkedin(db, document.id)
            if "success" in result and result["success"]:
                results["processed_now"] += 1
                if result["linkedin_urls_found"]:
                    results["linkedin_urls_found"].extend(result["linkedin_urls_found"])
                results["details"].append({
                    "document_id": document.id,
                    "filename": document.original_filename,
                    "success": True,
                    "urls_found": result["linkedin_urls_found"]
                })
            else:
                results["details"].append({
                    "document_id": document.id,
                    "filename": document.original_filename,
                    "success": False,
                    "error": result.get("error", "Unknown error")
                })
        
        # Remove duplicate URLs
        results["linkedin_urls_found"] = list(set(results["linkedin_urls_found"]))
        
        return results
    
    def update_linkedin_url(self, db: Session, user_id: int, linkedin_url: str) -> Dict[str, Any]:
        """Update LinkedIn URL for a user."""
        # Validate URL
        if not linkedin_url or 'linkedin.com' not in linkedin_url:
            return {"error": "Invalid LinkedIn URL"}
        
        # Get or create user profile
        profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
        if not profile:
            profile = UserProfile(user_id=user_id)
            db.add(profile)
        
        profile.linkedin_url = linkedin_url
        profile.updated_at = datetime.now()
        
        db.commit()
        
        return {
            "success": True,
            "user_id": user_id,
            "linkedin_url": linkedin_url,
            "message": "LinkedIn URL updated successfully"
        }